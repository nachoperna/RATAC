import re   # libreria de REGEX
import json
import os
import sys
from docx import Document
import pdfplumber
from pypdf import PdfReader

# ── CONSTANTES GLOBALES Y CONFIGURACIÓN ───────────────────────────────────────
MATERIAL_REMITIDO_ANTECEDENTES = "Material remitido - Antecedentes"
DESCRIPCION_MACROSCOPICA = "Descripción macroscópica"
DESCRIPCION_MICROSCOPICA = "Descripción microscópica"
DIAGNOSTICO_HISTOPATOLOGICO = "Diagnóstico histopatológico"
nombre_diag_actual = ""
CARGA_USUARIO = False

# ── EXPRESIONES REGULARES (REGEX) ─────────────────────────────────────────────
patron_referencias  = re.compile(r"^\s*(referencias?|anexos?)", re.I)
patronAntecedentes  = re.compile(r"material\s+remitido\s*-\s*antecedentes", re.I)
patronMacro         = re.compile(r"descripci[oó]n macrosc[oó]pica", re.I)
patronMicro         = re.compile(r"descripci[oó]n microsc[oó]pica", re.I)
patronDiag          = re.compile(r"diagn[oó]stico histopatol[oó]gico", re.I)
patronMorfologia    = re.compile(r"morfolog[ií]a", re.I)
patronInflamacion   = re.compile(r"inflamaci[oó]n", re.I)
patronDiagnostico   = re.compile(r"diagn[oó]stico", re.I)
patronTablaGrado    = re.compile(r"muestra|analizada", re.I)


# ==============================================================================
# MÓDULO 1: PROCESAMIENTO DE ARCHIVOS WORD (.docx)
# ==============================================================================

def procesarTablaGrado_docx(tabla):
    tabla_grado = []
    contenido = {
        "Caracteristica": "",
        "Muestra analizada": "",
        "Puntaje": ""
    }
    for row in tabla.rows[1:len(tabla.rows)-1]:
        contenido['Caracteristica'] = row.cells[0].text
        contenido['Muestra analizada'] = row.cells[1].text
        contenido['Puntaje'] = row.cells[2].text
        tabla_grado.append(contenido.copy())
        contenido = {
            "Caracteristica": "",
            "Muestra analizada": "",
            "Puntaje": ""
        }
    return tabla_grado

def datosPaciente_docx(tabla):
    datos = {}
    for row in tabla.rows:
        cells = row.cells
        if len(cells) >= 4:
            k1 = cells[0].text.strip()
            v1 = cells[1].text.strip()
            k2 = cells[2].text.strip()
            v2 = cells[3].text.strip()

            if "Raza" in k1 or "Edad" in k1:
                raza = re.search(r"^[^\d,-]+", v1)
                edad = re.search(r"\d+", v1)
                if raza: raza = raza.group().rstrip()
                if edad: edad = edad.group().rstrip()
                datos["Raza"] = raza
                datos["Edad"] = edad
            else:
                if k1:
                    if "Familia" in k1 and v1 == "": v1 = None
                    datos[k1] = v1

            if k2:
                if v2.startswith("-") and len(v2) == 5: v2 = f"01-01{v2}"
                if "Especie" in k2 and v2 == "": v2 = None
                datos[k2] = v2

    datos["Referencias mastocitomas"] = False
    return datos

def getImagenes_docx(doc, el):
    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    }
    rutas = []
    for dibujo in el.findall('.//w:drawing', namespaces=ns):
        for blip in dibujo.findall('.//a:blip', namespaces=ns):
            rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            if rId:
                imagen_part = doc.part.related_parts[rId]
                if imagen_part:
                    binario = imagen_part.blob
                    if len(binario) < 685000: continue
                    ruta = f"IMG_{rId}_{nombre_diag_actual}.png"
                    ruta = os.path.join("IMAGENES", ruta)
                    with open(ruta, "wb") as img:
                        img.write(binario)
                    rutas.append(ruta)
    return rutas

def matcheaCategoria(linea, seccion_actual, nueva_desc_micro, nombre_tabla_actual):
    if patronAntecedentes.search(linea):
        return True, MATERIAL_REMITIDO_ANTECEDENTES, nueva_desc_micro, nombre_tabla_actual
    elif patronMacro.search(linea):
        return True, DESCRIPCION_MACROSCOPICA, nueva_desc_micro, nombre_tabla_actual
    elif patronMicro.search(linea):
        return True, DESCRIPCION_MICROSCOPICA, True, nombre_tabla_actual
    elif patronDiag.search(linea):
        return True, DIAGNOSTICO_HISTOPATOLOGICO, nueva_desc_micro, nombre_tabla_actual
    elif patronMorfologia.search(linea) or patronInflamacion.search(linea):
        return True, seccion_actual, nueva_desc_micro, linea.strip()
    return False, seccion_actual, nueva_desc_micro, nombre_tabla_actual

def cargarDescMicro(descripciones_micro, bloque_actual):
    if bloque_actual['Diagnostico']['Descripcion'] in ("", "."):
        bloque_actual['Diagnostico']['Descripcion'] = None
    descripciones_micro.append(bloque_actual.copy())
    return {
        "Descripcion": "",
        "Diagnostico": {"Descripcion": "", "Imagenes": []},
        "Tabla de Grado": []
    }

def resultados(datos_paciente, secciones):
    resultado = {}
    resultado.update(datos_paciente)
    for k, v in secciones.items():
        if k != DESCRIPCION_MICROSCOPICA:
            resultado[k] = " ".join(v).strip() if v else "No encontrado"
        else:
            resultado[k] = v
    return resultado

def procesar_docx(archivo):
    ruta = archivo.path if hasattr(archivo, 'path') else archivo
    doc = Document(ruta)

    datos_paciente = {}
    secciones = {
        MATERIAL_REMITIDO_ANTECEDENTES: [],
        DESCRIPCION_MACROSCOPICA: [],
        DESCRIPCION_MICROSCOPICA: []
    }

    seccion_actual = None
    en_referencias = False
    cabecera_procesada = False
    map_parrafos = {p._element: p for p in doc.paragraphs}
    map_tablas = {t._element: t for t in doc.tables}
    nueva_desc_micro = False

    descripciones_micro = []
    bloque_actual = {
        "Descripcion": "",
        "Diagnostico": {"Descripcion": "", "Imagenes": []},
        "Tabla de Grado": []
    }
    nombre_tabla_actual = ""

    for el in doc.element.body:
        imagenes = getImagenes_docx(doc, el)
        if len(imagenes) > 0:
            bloque_actual["Diagnostico"]["Imagenes"].extend(imagenes)

        if not en_referencias:
            if el.tag.endswith("p"):
                p = map_parrafos.get(el)
                for r in p.runs:
                    linea = r.text
                    if not linea or patronDiagnostico.fullmatch(linea.strip()):
                        continue
                    if patron_referencias.search(linea):
                        en_referencias = True
                        if re.compile(r"referencia graduaci[oó]n mastocitomas", re.I).search(linea):
                            datos_paciente["Referencias mastocitomas"] = True
                        break

                    matchea_categoria, seccion_actual, nueva_desc_micro, nombre_tabla_actual = matcheaCategoria(linea, seccion_actual, nueva_desc_micro, nombre_tabla_actual)
                    if matchea_categoria: continue

                    if seccion_actual == DESCRIPCION_MICROSCOPICA:
                        if bloque_actual["Descripcion"] and nueva_desc_micro:
                            bloque_actual = cargarDescMicro(descripciones_micro, bloque_actual)
                        bloque_actual["Descripcion"] += linea.strip()
                    elif seccion_actual == DIAGNOSTICO_HISTOPATOLOGICO:
                        bloque_actual["Diagnostico"]["Descripcion"] += linea.strip()
                    else:
                        try:
                            secciones[seccion_actual].append(linea)
                        except Exception as e:
                            if CARGA_USUARIO:
                                print("Seccion desconocida", file=sys.stderr)
                                sys.exit(1)
                            with open("diagnosticos_mal_procesados.txt", "a", encoding="utf-8") as f:
                                f.write(f"{nombre_diag_actual}\n")
                            return ""
                    nueva_desc_micro = False

            elif el.tag.endswith("tbl"):
                tabla = map_tablas.get(el)
                if not cabecera_procesada:
                    datos_paciente = datosPaciente_docx(tabla)
                    cabecera_procesada = True
                else:
                    es_tabla_grado = (patronTablaGrado.match(tabla.rows[0].cells[1].text.strip()) if len(tabla.rows[0].cells) > 1 else False)
                    if es_tabla_grado:
                        bloque_actual['Tabla de Grado'] = procesarTablaGrado_docx(tabla)

    if bloque_actual["Descripcion"]:
        descripciones_micro.append(bloque_actual.copy())
    secciones[DESCRIPCION_MICROSCOPICA] = descripciones_micro

    return resultados(datos_paciente, secciones)


# ==============================================================================
# MÓDULO 2: PROCESAMIENTO DE ARCHIVOS PDF (.pdf)
# ==============================================================================

def limpiar_celda_pdf(celda):
    return (celda or "").replace('\n', ' ').strip()

def reordenar_datos_pdf(datos):
    orden_esperado = ["Protocolo", "Fecha", "Solicitante", "Técnica", "Propietario", "Especie", "Raza", "Edad", "Paciente", "Referencias mastocitomas"]
    resultado = {}
    for k in orden_esperado:
        val = datos.get(k, None)
        if val == "" or val == "-": val = None
        resultado[k] = val
    for k, v in datos.items():
        if k not in resultado:
            resultado[k] = None if (v == "" or v == "-") else v
    return resultado

def limpiar_diagnostico_pdf(desc):
    desc = re.sub(r"(?i)(Laura Denzoin|Médica Veterinaria|\bM\.?P\.?\b|Doctora en|Profesora Adjunta|Patología General|FCV-UNCPBA|Tandil\s*-|cov\.diagnostico|Dra\.?\s+Laura|Sof[íi]a Mart[íi]nez).*$", "", desc)
    return re.sub(r'\s+', ' ', desc).strip()

def procesar_raza_edad_pdf(val, datos):
    r_match = re.search(r"^[^\d]+", val)
    e_match = re.search(r"\d+", val)
    if r_match:
        raza_limpia = re.sub(r"(?i)años?|[-:]", "", r_match.group()).strip()
        if raza_limpia: datos["Raza"] = raza_limpia
    if e_match: datos["Edad"] = e_match.group().strip()
    return datos

def extraer_cabecera_regex_pdf(lineas_buffer):
    texto = " ".join(lineas_buffer)
    datos = {}
    m = re.search(r"Protocolo\s*:?\s*(\S+)", texto, re.I)
    if m: datos["Protocolo"] = m.group(1)
    m = re.search(r"Fecha\s*:?\s*([\d-]+)", texto, re.I)
    if m: datos["Fecha"] = m.group(1)
    m = re.search(r"Solicitante\s*:?\s*(.*?)\s*(?:T[ée]cnica|Propietario|Especie|$)", texto, re.I)
    if m: datos["Solicitante"] = m.group(1).strip(' -:,')
    m = re.search(r"T[ée]cnica\s*:?\s*(\S+)", texto, re.I)
    if m: datos["Técnica"] = m.group(1)
    m = re.search(r"Propietario\s*:?\s*(.*?)\s*(?:Especie|Raza|Paciente|Edad|$)", texto, re.I)
    if m: datos["Propietario"] = m.group(1).strip(' -:,')
    m = re.search(r"Especie\s*:?\s*(\S+)", texto, re.I)
    if m: datos["Especie"] = m.group(1)
    m = re.search(r"Raza[\s\-]*Edad\s*:?\s*(.*?)\s*(?:Paciente|$)", texto, re.I)
    if not m:
        m = re.search(r"Raza\s*:?\s*(.*?)\s*(?:Edad|Paciente|$)", texto, re.I)
    if m: datos = procesar_raza_edad_pdf(m.group(1), datos)
    m = re.search(r"Paciente\s*:?\s*(\S+)", texto, re.I)
    if m: datos["Paciente"] = m.group(1)
    datos["Referencias mastocitomas"] = bool(re.search(r"referencia graduaci[oó]n mastocitomas", texto, re.I))
    return datos

def procesarTablaGrado_pdf(tabla_filas):
    tabla_grado = []
    for row in tabla_filas[1:-1]:
        cells = [limpiar_celda_pdf(c) for c in row]
        if len(cells) >= 3:
            tabla_grado.append({
                "Caracteristica":   cells[0],
                "Muestra analizada": cells[1],
                "Puntaje":          cells[2]
            })
    return tabla_grado

def datosPaciente_pdf(tabla_filas):
    datos = {}
    for row in tabla_filas:
        cells = [limpiar_celda_pdf(c) for c in row]
        while len(cells) < 4: cells.append("")
        k1, v1, k2, v2 = cells[0], cells[1], cells[2], cells[3]
        raza_val = None
        if "Raza" in k1 or "Edad" in k1: raza_val = v1
        elif "Raza" in k2 or "Edad" in k2: 
            raza_val = v2
            k2 = "" 
        if raza_val: datos = procesar_raza_edad_pdf(raza_val, datos)
        if k1 and not ("Raza" in k1 or "Edad" in k1):
            if "Familia" in k1 and v1 == "": v1 = None
            datos[k1] = v1
        if k2:
            if v2.startswith("-") and len(v2) == 5: v2 = f"01-01{v2}"
            if "Especie" in k2 and v2 == "": v2 = None
            datos[k2] = v2
    if not datos.get("Raza") or datos.get("Raza") == "-":
        texto_tabla = " ".join([" ".join([c for c in row if c]) for row in tabla_filas])
        m = re.search(r"Raza[\s\-]*Edad\s*:?\s*(.*?)(?:Paciente|Especie|$)", texto_tabla, re.I)
        if m: datos = procesar_raza_edad_pdf(m.group(1), datos)
    datos["Referencias mastocitomas"] = False
    return datos

def getImagenes_pdf(reader, page_num):
    rutas = []
    try:
        page = reader.pages[page_num]
        for idx, image_obj in enumerate(page.images):
            binario = image_obj.data
            if len(binario) < 685_000: continue
            nombre_img = f"IMG_p{page_num}_{idx}_{nombre_diag_actual}.png"
            ruta = os.path.join("IMAGENES", nombre_img)
            with open(ruta, "wb") as img:
                img.write(binario)
            rutas.append(ruta)
    except Exception:
        pass
    return rutas

def agrupar_palabras_en_lineas(words, tolerancia=3):
    if not words: return []
    words_sorted = sorted(words, key=lambda w: (w["top"], w["x0"]))
    lineas = []
    linea_actual = [words_sorted[0]]
    top_ref = words_sorted[0]["top"]
    for word in words_sorted[1:]:
        if abs(word["top"] - top_ref) <= tolerancia:
            linea_actual.append(word)
        else:
            lineas.append(sorted(linea_actual, key=lambda w: w["x0"]))
            linea_actual = [word]
            top_ref = word["top"]
    if linea_actual:
        lineas.append(sorted(linea_actual, key=lambda w: w["x0"]))
    return lineas

def procesar_pdf(ruta):
    datos_paciente = {}
    secciones = {
        MATERIAL_REMITIDO_ANTECEDENTES: [],
        DESCRIPCION_MACROSCOPICA: [],
        DESCRIPCION_MICROSCOPICA: []
    }
    seccion_actual = None
    en_referencias = False
    cabecera_procesada = False
    nueva_desc_micro = False
    cabecera_buffer = []
    en_tabla_grado_texto = False
    descripciones_micro = []
    bloque_actual = {
        "Descripcion": "",
        "Diagnostico": {"Descripcion": "", "Imagenes": []},
        "Tabla de Grado": []
    }
    nombre_tabla_actual = ""
    reader = PdfReader(ruta)

    with pdfplumber.open(ruta) as pdf:
        for page_num, page in enumerate(pdf.pages):
            if en_referencias: break
            imagenes = getImagenes_pdf(reader, page_num)
            if imagenes:
                bloque_actual["Diagnostico"]["Imagenes"].extend(imagenes)
            tablas_obj = page.find_tables()
            tabla_bboxes = [t.bbox for t in tablas_obj]
            todas_las_palabras = page.extract_words(x_tolerance=1.5, keep_blank_chars=True)
            palabras_fuera = [w for w in todas_las_palabras if not any(w["x0"] >= bx[0] and w["x1"] <= bx[2] and w["top"] >= bx[1] and w["bottom"] <= bx[3] for bx in tabla_bboxes)]
            lineas_texto = agrupar_palabras_en_lineas(palabras_fuera)
            elementos = [(l[0]["top"], "texto", " ".join(w["text"] for w in l)) for l in lineas_texto]
            for t_obj in tablas_obj:
                elementos.append((t_obj.bbox[1], "tabla", t_obj))
            elementos.sort(key=lambda e: e[0])

            for _, tipo, contenido in elementos:
                if en_referencias: break
                if tipo == "texto":
                    linea = contenido
                    if not linea or patronDiagnostico.fullmatch(linea.strip()):
                        continue
                    if patron_referencias.search(linea):
                        en_referencias = True
                        if re.compile(r"referencia graduaci[oó]n mastocitomas", re.I).search(linea):
                            datos_paciente["Referencias mastocitomas"] = True
                        break
                    matchea, nueva_seccion, nueva_desc_micro, nombre_tabla_actual = \
                        matcheaCategoria(linea, seccion_actual, nueva_desc_micro, nombre_tabla_actual)
                    if matchea:
                        en_tabla_grado_texto = False
                        if seccion_actual is None:
                            if not cabecera_procesada and cabecera_buffer:
                                datos = extraer_cabecera_regex_pdf(cabecera_buffer)
                                if datos:
                                    datos_paciente.update(datos)
                                    cabecera_procesada = True
                            datos_paciente = reordenar_datos_pdf(datos_paciente)
                        seccion_actual = nueva_seccion
                        continue
                    if seccion_actual is None:
                        cabecera_buffer.append(linea)
                        continue
                    if seccion_actual == DESCRIPCION_MICROSCOPICA:
                        if bloque_actual["Descripcion"] and nueva_desc_micro:
                            bloque_actual = cargarDescMicro(descripciones_micro, bloque_actual)
                        if re.search(r"Características.*Muestra analizada|Muestra analizada.*Puntaje", linea, re.I):
                            en_tabla_grado_texto = True
                            continue
                        if en_tabla_grado_texto:
                            m = re.search(r"^(Diferenciaci[oó]n|Mitosis|Necrosis)\s+([\w%]+)\s+(\d+)", linea.strip(), re.I)
                            if m:
                                bloque_actual["Tabla de Grado"].append({
                                    "Caracteristica": m.group(1).capitalize(),
                                    "Muestra analizada": m.group(2),
                                    "Puntaje": m.group(3)
                                })
                            elif re.search(r"total\s*\d+", linea, re.I):
                                en_tabla_grado_texto = False
                            continue 
                        bloque_actual["Descripcion"] += linea.strip() + " "
                    elif seccion_actual == DIAGNOSTICO_HISTOPATOLOGICO:
                        bloque_actual["Diagnostico"]["Descripcion"] += linea.strip() + " "
                    else:
                        secciones[seccion_actual].append(linea)
                    nueva_desc_micro = False
                elif tipo == "tabla":
                    tabla_filas = contenido.extract()
                    if not tabla_filas: continue
                    if seccion_actual is None:
                        for row in tabla_filas:
                            cabecera_buffer.append(" ".join([limpiar_celda_pdf(c) for c in row if c]))
                    if not cabecera_procesada:
                        datos_paciente.update(datosPaciente_pdf(tabla_filas))
                        if datos_paciente: cabecera_procesada = True
                    else:
                        primera_fila = tabla_filas[0] if tabla_filas else []
                        es_tabla_grado = (len(primera_fila) > 1 and primera_fila[1] and patronTablaGrado.match(limpiar_celda_pdf(primera_fila[1])))
                        if es_tabla_grado:
                            bloque_actual["Tabla de Grado"] = procesarTablaGrado_pdf(tabla_filas)

    if bloque_actual["Descripcion"]:
        bloque_actual["Descripcion"] = re.sub(r'\s+', ' ', bloque_actual["Descripcion"]).strip()
        bloque_actual["Diagnostico"]["Descripcion"] = limpiar_diagnostico_pdf(bloque_actual["Diagnostico"]["Descripcion"])
        descripciones_micro.append(bloque_actual.copy())
    secciones[DESCRIPCION_MICROSCOPICA] = descripciones_micro
    return resultados(datos_paciente, secciones)


# ==============================================================================
# MÓDULO 3: ENRUTADOR PRINCIPAL (Maneja masivo y por parámetro individual)
# ==============================================================================

if __name__ == "__main__":
    param = sys.argv[1] if len(sys.argv) > 1 else None
    os.makedirs("JSONS/", exist_ok=True)
    os.makedirs("IMAGENES/", exist_ok=True)

    if not param:
        # Modo Masivo (Escanea la carpeta de Histopatología)
        ruta = "./Histopatología/"
        if os.path.exists(ruta):
            with os.scandir(ruta) as archivos:
                for archivo in archivos:
                    nombre = os.path.basename(archivo.name)
                    nombre_diag_actual = nombre
                    nombre, ext = os.path.splitext(nombre)
                    
                    data = None
                    if ext.lower() == ".pdf":
                        print("Procesando PDF:", nombre)
                        data = procesar_pdf(archivo.path)
                    elif ext.lower() == ".docx":
                        print("Procesando DOCX:", nombre)
                        data = procesar_docx(archivo.path)
                    else:
                        continue

                    if data:
                        ruta_salida = os.path.join("JSONS/", nombre + ".json")
                        with open(ruta_salida, "w", encoding="utf-8") as f:
                            json.dump(data, f, indent=4, ensure_ascii=False)
        else:
            print(f"Carpeta {ruta} no encontrada.")
    else:
        # Modo Individual (Llamado desde Go / paciente_handler.go)
        CARGA_USUARIO = True
        ruta_archivo = param
        
        # Mantenemos la lógica de captura de nombre que tenía tu diag_to_json original
        if len(sys.argv) > 2:
            nombre_diag_actual = sys.argv[2]
        else:
            nombre_diag_actual = os.path.basename(ruta_archivo)
            
        nombre, ext = os.path.splitext(os.path.basename(ruta_archivo))

        data = None
        if ext.lower() == ".pdf":
            data = procesar_pdf(ruta_archivo)
        elif ext.lower() == ".docx":
            data = procesar_docx(ruta_archivo)

        if data:
            ruta_salida = os.path.join("JSONS/", nombre + ".json")
            with open(ruta_salida, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=4, ensure_ascii=False)
            
            # Imprimimos el JSON en consola (stdout) para que Go lo reciba si lo necesita
            print(json.dumps(data, indent=4, ensure_ascii=False))