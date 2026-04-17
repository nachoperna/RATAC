import re   # libreria de REGEX
import json
import os
from docx import Document
# from pdf2docx import Converter

# CONSTANTES
MATERIAL_REMITIDO_ANTECEDENTES = "Material remitido - Antecedentes"
DESCRIPCION_MACROSCOPICA = "Descripción macroscópica"
DESCRIPCION_MICROSCOPICA = "Descripción microscópica"
DIAGNOSTICO_HISTOPATOLOGICO = "Diagnóstico histopatológico"
nombre_diag_actual = ""

# regexs que nos indican los cambios de seccion o elementos a evitar
patron_referencias = re.compile(r"^\s*(referencias?|anexos?)", re.I)   # texto que comience con la palabra referencias o anexo, en plural o singular
patronAntecedentes = re.compile(r"material\s+remitido\s*-\s*antecedentes", re.I)
patronMacro = re.compile(r"descripci[oó]n macrosc[oó]pica", re.I)
patronMicro = re.compile(r"descripci[oó]n microsc[oó]pica", re.I)
patronDiag = re.compile(r"diagn[oó]stico histopatol[oó]gico", re.I)
patronMorfologia = re.compile(r"morfolog[ií]a", re.I)
patronInflamacion = re.compile(r"inflamaci[oó]n", re.I)
patronDiagnostico = re.compile(r"diagn[oó]stico", re.I)
patronTablaGrado = re.compile(r"muestra|analizada", re.I)

def procesarTabla(tabla, nombre_diag):
    try:
        filas = []
        categorias = []

        # iteramos sobre la primer fila de la tabla y obtenemos sus columnas
        
        for cell in tabla.rows[0].cells:
            if cell.text.strip() == "":
                continue
            categorias.append(cell.text)

        for row in tabla.rows[1:]:  # iteramos a partir de la segunda fila
            # FORMATO DE SALIDA: {contenido de fila: tipo de categoria a la que corresponde}
            fila = "{" + row.cells[0].text
            # iteramos a partir de la segunad celda para obtener el valor de las columnas
            for j, cell in enumerate(row.cells[1:]):
                if cell.text.strip():   # si en la celda hay un string (normalmente X) entonces obtenemos el tipo de categoria a la que pertence la fila usando el indice sobre el que iteramos
                    fila += ": " + categorias[j] + "}"
            filas.append(fila)
        return ", ".join(filas)
    except Exception as e:
        print(f"!-ERROR: {e}")
        with open("diagnosticos_mal_procesados.txt", "a", encoding="utf-8") as f:
            f.write(f"{nombre_diag}\n")
        return ""

def procesarTablaGrado(tabla):
    tabla_grado = []
    contenido = {
        "Caracteristica": "",
        "Muestra analizada": "",
        "Puntaje": ""
    }

    for row in tabla.rows[1:len(tabla.rows)-1]:  # iteramos a partir de la segunda fila
        # FORMATO DE SALIDA: {contenido de fila: tipo de categoria a la que corresponde}
        contenido['Caracteristica'] = row.cells[0].text
        contenido['Muestra analizada'] = row.cells[1].text
        contenido['Puntaje'] = row.cells[2].text
        tabla_grado.append(contenido.copy())
        contenido = {
            "Caracteristica": "",
            "Muestra analizada": "",
            "Puntaje": ""
        }

    return tabla_grado

def datosPaciente(tabla):
    datos = {}

    for row in tabla.rows:
        cells = row.cells
        if len(cells) >= 4:
            k1 = cells[0].text.strip()
            v1 = cells[1].text.strip()
            k2 = cells[2].text.strip()
            v2 = cells[3].text.strip()

            if "Raza" in k1 or "Edad" in k1:
                # nos quedamos con todo lo que no sea digitos, comas o guiones
                raza = re.search(r"^[^\d,-]+", v1)
                # nos quedamos solo con los digitos
                edad = re.search(r"\d+", v1)

                if raza:
                    # rstrip elimina todos los 2 o mas espacios juntos luego de todos los caracteres
                    raza = raza.group().rstrip()
                if edad:
                    edad = edad.group().rstrip()
                datos["Raza"] = raza
                datos["Edad"] = edad
            else:
                if k1:
                    datos[k1] = v1

            if k2:
                if v2.startswith("-") and len(v2) == 5:
                    v2 = f"01-01{v2}"
                datos[k2] = v2

    datos["Referencias mastocitomas"] = False
    return datos

def getImagenes(doc, el):
    ns = {  # aparentemente hay que hacer esto para obtener bien cada imagen del doc
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    }
    rutas = []
    for dibujo in el.findall('.//w:drawing', namespaces=ns):
        # la etiqueta blip que contiene el ID de la imagen
        for blip in dibujo.findall('.//a:blip', namespaces=ns):
            # obtenemos el ID de la imagen
            rId = blip.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")

            if rId:
                imagen_part = doc.part.related_parts[rId]
                if imagen_part:
                    binario = imagen_part.blob  # blob nos da el binario
                    # si el binario es menor a 685k bits o 85Kb entonces seguramente sea una foto de una firma digital del doc y la evitamos directamente
                    if len(binario) < 685000:
                        continue
                    # todas las imagenes de los diagnosticos ya estan en formato png
                    ruta = f"IMG_{rId}_{nombre_diag_actual}.png"
                    ruta = os.path.join("./IMAGENES", ruta)
                    with open(ruta, "wb") as img:   # wb = write binary
                        img.write(binario)
                    rutas.append(ruta)
    # retornamos todas las rutas (relativas) que luego son asignadas a cada parte ed los diagnosticos
    return rutas

def tablaProcesable(contenido, patron):
    return contenido or patron.match(contenido.strip())

def matcheaCategoria(linea, seccion_actual, nueva_desc_micro, nombre_tabla_actual):
    # seteamos la seccion actual
    if patronAntecedentes.search(linea):
        return True, MATERIAL_REMITIDO_ANTECEDENTES, nueva_desc_micro, nombre_tabla_actual
    elif patronMacro.search(linea):
        return True, DESCRIPCION_MACROSCOPICA, nueva_desc_micro, nombre_tabla_actual
    elif patronMicro.search(linea):
        # levantamos la flag que indica que debemos guardar el bloque actual de descripcion que tenemos hasta el momento
        return True, DESCRIPCION_MICROSCOPICA, True, nombre_tabla_actual
    elif patronDiag.search(linea):
        return True, DIAGNOSTICO_HISTOPATOLOGICO, nueva_desc_micro, nombre_tabla_actual
    elif patronMorfologia.search(linea) or patronInflamacion.search(linea):     # signfica que estamos por procesar una tabla que nos sirve
        return True, seccion_actual, nueva_desc_micro, linea.strip()
    
    return False, seccion_actual, nueva_desc_micro, nombre_tabla_actual

def cargarDescMicro(descripciones_micro, bloque_actual):
    # no se puede hacer un simple append(bloque_actual) porque las variables son mutables entonces tener que hacer un copy de cada estructura
    descripciones_micro.append(bloque_actual.copy())
    return {   # reiniciamos la estructura
        "Descripcion": "",
        "Diagnostico": {
            "Descripcion": "",
            "Imagenes": []
        },
        "Tabla de Grado": []
    }

def resultados(datos_paciente, secciones):
    resultado = {}  # donde juntamos todas las estructuras para luego formar el json con dump()
    resultado.update(datos_paciente)

    for k, v in secciones.items():
        if k != DESCRIPCION_MICROSCOPICA:
            resultado[k] = " ".join(v).strip() if v else "No encontrado"
        else:
            resultado[k] = v
    return resultado

def procesar_docx(ruta):
    doc = Document(ruta)

    datos_paciente = {}  # guardamos la primer tabla con los datos basicos
    secciones = {   # cada seccion va a guardar datos y/o conjunto de datos
        MATERIAL_REMITIDO_ANTECEDENTES: [],
        DESCRIPCION_MACROSCOPICA: [],
        DESCRIPCION_MICROSCOPICA: []
    }

    seccion_actual = None
    en_referencias = False
    cabecera_procesada = False

    # mapeamos todos los parrafos y tablas del doc para dsp al iterar uno por uno en el XML directo ir a buscar ese elemento procesado a estos mapas
    map_parrafos = {p._element: p for p in doc.paragraphs}
    map_tablas = {t._element: t for t in doc.tables}

    # flag que nos indica cuando dar de alta una nueva descripcion microscopica en secciones e iniciar otra
    nueva_desc_micro = False

    descripciones_micro = []
    bloque_actual = {   # solo utilizado para las descripciones microscopicas
        "Descripcion": "",
        "Diagnostico": {    # cada descripcion microscopica tiene su diagnostico asociado
            "Descripcion": "",
            "Imagenes": []
        },
        "Tabla de Grado": []
    }
    nombre_tabla_actual = ""    # seguramente solo morfologia o inflamacion

    for el in doc.element.body:
        imagenes = getImagenes(doc, el)
        if len(imagenes) > 0:
            bloque_actual["Diagnostico"]["Imagenes"].extend(imagenes)

        if not en_referencias:  # ignoramos todo el contenido de referencias o anexos

            if el.tag.endswith("p"):    # estamos ante un parrafo
                p = map_parrafos.get(el)    # obtenemos el parrafo pre procesado
                
                for r in p.runs:    # iteramos por cada linea del parrafo
                    linea = r.text

                    # si la linea esta vacia o va a empezar una tabla de diagnostico salteamos esta iteracion
                    if not linea or patronDiagnostico.fullmatch(linea.strip()):
                        continue

                    # si nos encontramos con una liena que solo diga Referencias cortamos el procesado del parrafo actual
                    if patron_referencias.search(linea):
                        en_referencias = True
                        if re.compile(r"referencia graduaci[oó]n mastocitomas", re.I).search(linea):
                            datos_paciente["Referencias mastocitomas"] = True
                        break

                    matchea_categoria, seccion_actual, nueva_desc_micro, nombre_tabla_actual = matcheaCategoria(linea, seccion_actual, nueva_desc_micro, nombre_tabla_actual)
                    if matchea_categoria:
                        continue

                    if seccion_actual == DESCRIPCION_MICROSCOPICA:
                        # si el bloque actual tiene contenido y debemos guardarlo
                        if bloque_actual["Descripcion"] and nueva_desc_micro:
                            bloque_actual = cargarDescMicro(descripciones_micro, bloque_actual)

                        bloque_actual["Descripcion"] += linea.strip()
                    elif seccion_actual == DIAGNOSTICO_HISTOPATOLOGICO:
                        bloque_actual["Diagnostico"]["Descripcion"] += linea.strip()
                    else:
                        try:
                            secciones[seccion_actual].append(linea)
                        except Exception as e:
                            print(f"ERROR: {e}")
                            with open("diagnosticos_mal_procesados.txt", "a", encoding="utf-8") as f:
                                f.write(f"{nombre_diag_actual}\n")
                            return ""

                    # volvemos al estado original para seguir completando el bloque actual
                    nueva_desc_micro = False

            elif el.tag.endswith("tbl"):    # estamos ante una tabla
                tabla = map_tablas.get(el)  # obtenemos la tabla pre procesada
                if not cabecera_procesada:  # procesamos la primer tabla que tiene los datos basicos del paciente
                    datos_paciente = datosPaciente(tabla)
                    cabecera_procesada = True
                # solo procesamos tablas NO RALAS (por ejemplo dejamos afuera la de diagnostico porque dice lo mismo que la descripcion del diagnostico)
                else:
                    # contenido_primera_celda = tabla.rows[0].cells[0].text
                    es_tabla_grado = (patronTablaGrado.match(tabla.rows[0].cells[1].text.strip()) if len(tabla.rows[0].cells) > 1 else False)
                    # tabla_procesable = (not contenido_primera_celda) or es_tabla_grado
                    if es_tabla_grado:
                        bloque_actual['Tabla de Grado'] = procesarTablaGrado(tabla)
                    # if tabla_procesable:
                    #     texto_tabla = procesarTabla(tabla, os.path.basename(ruta))
                    #
                    #     # si estamos en una seccion de descripcion micro o de diagnostico, agregamos la tabla al diagnostico asociado
                    #     if seccion_actual in [DESCRIPCION_MICROSCOPICA, DIAGNOSTICO_HISTOPATOLOGICO]:
                    #         bloque_actual["Diagnostico"]["Tablas"].append({
                    #             "Nombre": nombre_tabla_actual,
                    #             "Contenido": texto_tabla
                    #         })
                    #         nombre_tabla_actual = ""  # limpiamos para la próxima tabla
                    #     else:
                    #         try:
                    #             secciones[seccion_actual].append(texto_tabla)
                    #         except Exception as e:
                    #             print(f"ERRO : {e}")
                    #             with open("diagnosticos_mal_procesados.txt", "a", encoding="utf-8") as f:
                    #                 f.write(f"{nombre_diag_actual}\n")
                    #             return ""

    # agregamos el ultimo contenido de descripcion microscopica del documento
    if bloque_actual["Descripcion"]:
        descripciones_micro.append(bloque_actual.copy())

    secciones[DESCRIPCION_MICROSCOPICA] = descripciones_micro

    return resultados(datos_paciente, secciones)


# def convertDocx(ruta):
#     cv = Converter(ruta)
#     cv.convert(ruta + ".docx")
#     cv.close()
#     os.remove(ruta)

if __name__ == "__main__":
    ruta = "../Histopatología/"

    if not os.path.exists("JSONS/"):
        os.makedirs("JSONS/")
    if not os.path.exists("IMAGENES/"):
        os.makedirs("IMAGENES/")

    with os.scandir(ruta) as archivos:
        for archivo in archivos:
            nombre = os.path.basename(archivo.name)
            nombre_diag_actual = nombre
            nombre, ext = os.path.splitext(nombre)
            
            if ext == ".pdf":
                # convertDocx(os.path.abspath(archivo)) 
                continue

            print("Procesando: ", nombre)
            data = procesar_docx(archivo)

            if data:
                nombre = os.path.join("JSONS/", nombre)
                nombre += ".json"
                with open(nombre, "w", encoding="utf-8") as f:
                    json.dump(data, f, indent=4, ensure_ascii=False)
